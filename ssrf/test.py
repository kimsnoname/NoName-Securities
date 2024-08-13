from docx import Document
from docx.shared import Inches
import os

# Load the document
doc_path = "/mnt/data/인증정보 탈취.docx"
doc = Document(doc_path)

# Create a new document
new_doc = Document()

# Step 1 to Step 10
steps_1_to_10 = [
    "step 1. SSRF 취약점 확인",
    "step 2. ImageURL을 http://ifconfig.io 변경하여 전송",
    "step 3. 응답을 base64로 복호화",
    "step 4. HTML 소스코드 확인 후 파일 변환",
    "step 5. HTML 열람하여 BACKEND 요청 확인 및 AWS 서울의 EC2 인 것을 확인",
    "step 6. ImageURL을 http://169.254.169.254/latest/meta-data 입력",
    "step 7. meta-data 응답 확인",
    "step 8. 이후 반복적인 request를 효율화 하기 위하여 flask 작성",
    "step 9. http://169.254.169.254/latest/meta-data/iam/security-credentials/EC2_Access/",
    "step 10. AccessKeyId SecretAccessKey Token 탈취"
]

for step in steps_1_to_10:
    new_doc.add_paragraph(step)

# Add the first image
new_doc.add_picture("/mnt/data/image-0.png", width=Inches(6))

# Step 11 to Step 26
steps_11_to_26 = [
    "step 11. CMD 오픈 및 aws configure 입력",
    "step 12. AccessKeyId SecretAccessKey region 입력",
    "step 13. aws configure set aws_session_token '토큰’ 입력",
    "step 14. aws sdk로 adminKey 입력",
    "aws ec2 create-key-pair --key-name Yesnamekey --query “KeyMaterial” --output text >  YesnameKey.pem",
    "step 15. 자동화 프로그램 실행",
    "VPCSUBNETIGWRTSGEC2 생성 확인",
    "step 16. AWS SDK에 접속하여 iam 계정 추가 및 권한 취득",
    "step 17. AWS WEB 접속",
    "step 18. AWS EC2 접속",
    "step 19. EC2 instancce선택 후 연결",
    "step 20. RDP 클라이언트 선택",
    "step 21. 암호 가져오기 선택",
    "step 22. 프라이빗 키 파일을 업로드",
    "step 23. 암호 해독",
    "step 24. PUBLIC IP 사용자 이름 암호 확인 후 RDP 접속",
    "step 25. 코인 채굴 확인",
    "step 26. 코인 입금 확인"
]

for step in steps_11_to_26:
    new_doc.add_paragraph(step)

# Add the second image
new_doc.add_picture("/mnt/data/image-1.png", width=Inches(6))

# Step 27 to Step 34
steps_27_to_34 = [
    "step 27. http://169.254.169.254/latest/meta-data/instance-id/ 에서 instance id 확인",
    "step 28. Backend instance id에 SSM으로 연결",
    "step 29. authorized_keys 경로 이동",
    "Sudo su 입력 root 권한 획득 Cd /home/ubuntu/.ssh 로 이동 후 새 adminkey 생성",
    "step 30. 새 키페어 생성",
    "step 31. Adminkey.pub 내용 확인 및 삽입",
    "step 32. 대칭키 저장",
    "step 33. SSH 로그인",
    "step 34. 침입 흔적 삭제"
]

for step in steps_27_to_34:
    new_doc.add_paragraph(step)

# Add remaining images
for i in range(2, 11):
    image_path = f"/mnt/data/image-{i}.png"
    if os.path.exists(image_path):
        new_doc.add_picture(image_path, width=Inches(6))

# Save the new document
new_doc.save("/mnt/data/정리된_인증정보_탈취.docx")

import ace_tools as tools; tools.display_file("정리된_인증정보_탈취.docx")
